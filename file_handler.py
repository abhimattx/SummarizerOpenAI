import os
import io
import fitz  # PyMuPDF
import docx
import tempfile
import base64
from PIL import Image

def read_pdf(file):
    """Extract text from a PDF file."""
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def read_pdf_with_metadata(file):
    """Extract text and metadata from a PDF file."""
    metadata = {}
    text = ""
    file_bytes = file.read()
    
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        # Extract metadata
        metadata = {
            "Page Count": doc.page_count,
            "Title": doc.metadata.get("title", ""),
            "Author": doc.metadata.get("author", ""),
            "Subject": doc.metadata.get("subject", ""),
            "Creator": doc.metadata.get("creator", ""),
            "Producer": doc.metadata.get("producer", ""),
        }
        
        # Extract text
        for page in doc:
            text += page.get_text()
            
    return metadata, text

def read_docx(file):
    """Extract text from a DOCX file."""
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_docx_with_metadata(file):
    """Extract text and metadata from a DOCX file."""
    # Save a temp copy of the file to access metadata
    file_bytes = file.read()
    file.seek(0)  # Reset file pointer
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file.write(file_bytes)
        temp_path = temp_file.name
    
    try:
        doc = docx.Document(temp_path)
        
        # Extract text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)
        
        # Extract basic metadata
        metadata = {
            "Page Count": len(doc.paragraphs),
            "Word Count": len(text.split()),
            "Characters": len(text)
        }
        
        # Try to extract core properties if available
        try:
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "Title": core_props.title or "",
                    "Author": core_props.author or "",
                    "Created": core_props.created or "",
                    "Modified": core_props.modified or "",
                })
        except:
            pass  # Some properties might not be available
            
        return metadata, text
    finally:
        # Clean up temp file
        try:
            os.remove(temp_path)
        except:
            pass

def read_txt(file):
    """Read text from a TXT file."""
    return file.read().decode('utf-8')

def read_txt_with_metadata(file):
    """Extract text and basic metadata from a TXT file."""
    text = file.read().decode('utf-8')
    
    # Calculate basic metadata
    metadata = {
        "Line Count": text.count('\n') + 1,
        "Word Count": len(text.split()),
        "Characters": len(text)
    }
    
    return metadata, text

def save_uploaded_file(uploaded_file):
    """Save the uploaded file temporarily and return its path."""
    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
    os.makedirs(upload_dir, exist_ok=True)
    
    temp_path = os.path.join(upload_dir, uploaded_file.name)
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.getvalue())
    
    return temp_path

def extract_text_from_file(uploaded_file):
    """Extract text from various file formats."""
    if uploaded_file is None:
        return None
    
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return read_pdf(uploaded_file)
        elif file_extension == '.docx':
            return read_docx(uploaded_file)
        elif file_extension == '.txt':
            return read_txt(uploaded_file)
        else:
            return None
    except Exception as e:
        raise Exception(f"Error processing file: {str(e)}")

def extract_text_from_file_with_metadata(uploaded_file):
    """Extract text and metadata from various file formats."""
    if uploaded_file is None:
        return None, None
    
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return read_pdf_with_metadata(uploaded_file)
        elif file_extension == '.docx':
            return read_docx_with_metadata(uploaded_file)
        elif file_extension == '.txt':
            return read_txt_with_metadata(uploaded_file)
        else:
            return None, None
    except Exception as e:
        raise Exception(f"Error processing file metadata: {str(e)}")

def get_file_preview(uploaded_file):
    """Generate a preview image for supported file formats."""
    if uploaded_file is None:
        return None
    
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    
    try:
        if file_extension == '.pdf':
            # Get the first page of the PDF as an image
            file_bytes = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer
            
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                if len(doc) > 0:
                    page = doc[0]  # First page
                    pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                    img_bytes = pix.tobytes("png")
                    return Image.open(io.BytesIO(img_bytes))
        
        # For other file types, return None (will use default icons)
        return None
            
    except Exception as e:
        print(f"Preview generation error: {str(e)}")
        return None